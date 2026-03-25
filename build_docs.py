"""
Generate .docx documentation from markdown files with ORNL / Pellissippi State theme.

Colors:
  Title blue:   #0070C0
  Gold accent:  #B48C32
  Heading navy: #142850
  Body text:    #333333
  Meta gray:    #64646E
  Code bg:      light gray (shading)
  Table header: #142850 bg, white text
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re

OUT = "C:/Users/Lenovo/OneDrive/Desktop/Documentation"

BLUE    = RGBColor(0x00, 0x70, 0xC0)
GOLD    = RGBColor(0xB4, 0x8C, 0x32)
NAVY    = RGBColor(0x14, 0x28, 0x50)
BODY    = RGBColor(0x33, 0x33, 0x33)
GRAY    = RGBColor(0x64, 0x64, 0x6E)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG = "E8E8E8"
TBL_HDR = "142850"


def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BODY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, size in [(1, 16), (2, 13), (3, 11)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = NAVY
        h.font.bold = True
        h.paragraph_format.space_before = Pt(18 if level == 1 else 12)

    return doc


def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.color.rgb = BLUE
    run.font.bold = True
    p.paragraph_format.space_after = Pt(4)
    # Gold underline bar
    bar = doc.add_paragraph()
    r = bar.add_run("━" * 60)
    r.font.color.rgb = GOLD
    r.font.size = Pt(10)
    bar.paragraph_format.space_after = Pt(12)


def add_code(doc, text):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}"/>')
        p.paragraph_format.element.get_or_add_pPr().append(shading)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TBL_HDR}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)


def md_to_docx(md_path, docx_path):
    """Convert a markdown file to a themed .docx document."""
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = new_doc()
    i = 0
    in_code = False
    code_buf = []
    in_table = False
    table_headers = []
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Code block toggle
        if line.startswith("```"):
            if in_code:
                add_code(doc, "\n".join(code_buf))
                code_buf = []
                in_code = False
            else:
                # Flush any pending table
                if in_table:
                    add_table(doc, table_headers, table_rows)
                    in_table = False
                    table_headers = []
                    table_rows = []
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Table row
        if "|" in line and line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # Skip separator rows (|---|---|)
            if all(set(c) <= {"-", ":", " "} for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            add_table(doc, table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []

        # Horizontal rule
        if line.strip() == "---":
            i += 1
            continue

        # Headings
        if line.startswith("# ") and not line.startswith("## "):
            add_title(doc, line[2:].strip())
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # Bullet points
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            # Handle bold at start of bullet
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.+)", line)
        if m:
            text = m.group(2)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph (handle inline bold/code)
        text = line.strip()
        p = doc.add_paragraph()
        # Split on bold markers and inline code
        parts = re.split(r"(\*\*.*?\*\*|`[^`]+`)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            else:
                p.add_run(part)
        i += 1

    # Flush remaining table
    if in_table:
        add_table(doc, table_headers, table_rows)

    doc.save(docx_path)
    print(f"  Saved: {docx_path}")


def main():
    import os

    files = {
        "architecture.md": "CTI_Architecture.docx",
        "adding_a_source.md": "Adding_a_Source.docx",
        "feed_configs.md": "Feed_Configurations.docx",
        "example_abuseipdb.md": "Example_AbuseIPDB.docx",
    }

    for md_name, docx_name in files.items():
        md_path = os.path.join(OUT, md_name)
        docx_path = os.path.join(OUT, docx_name)
        if os.path.exists(md_path):
            md_to_docx(md_path, docx_path)
        else:
            print(f"  Skipped (not found): {md_path}")

    print("\nDone.")


if __name__ == "__main__":
    main()
